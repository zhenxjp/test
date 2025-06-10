import docx
from docx.text.run import Run
from docx.text.paragraph import Paragraph

from lxml import etree
from xxx_docx import *

import sys





#############################################################################
# 测试：打印每部分属性
def doc_do_text_att(doc_path:str):
    doc = docx.Document(doc_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run_is_delete_in(run,'.//w:dstrike'):
                print(['double delete'])

            if run_is_delete_in(run,'.//w:strike'):
                print(['delete'])


            if run_is_highlight(run):
                print(['highlight'])

            print(f'{run.text}\n')

            print(run_att_to_str(run))
            print('==============================================')
#############################################################################



# 获取修订前，修订后的信息
def doc_recovery(doc_path:str):


    doc = docx.Document(doc_path)
    return doc_do_paragraph(doc.paragraphs)
    
#############################################################################
# 打印表格
def doc_read_table(doc_path:str):
    doc = docx.Document(doc_path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                old_text,new_text  = doc_do_paragraph(cell.paragraphs,end = '')
                print(f'{old_text}  ',end='')
            print('\n')
    print('------------------------------------')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                old_text,new_text  = doc_do_paragraph(cell.paragraphs,end = '')
                print(f'{new_text}  ',end='')
            print('\n')
    print('====================================')
#############################################################################

def test_doc_text():
    doc_path = 'xxx.docx'  # 替换为实际文档路径
    #doc_path = 'highlight.docx'  # 替换为实际文档路径
    doc_do_text_att(doc_path)

    old_text, new_text = doc_recovery(doc_path)
    print('==================old============================')
    print(old_text)
    # 写入 old.txt 文件
    with open('old.txt', 'w', encoding='utf-8') as file:
        file.write(old_text)
    print('==================new============================')
    print(new_text)
    # 写入 new.txt 文件
    with open('new.txt', 'w', encoding='utf-8') as file:
        file.write(new_text)

def test_doc_tables():
    doc_path = 'table.docx'  # 替换为实际文档路径
    doc_read_table(doc_path)
    doc_do_text_att(doc_path)

if __name__ == "__main__":
    test_doc_tables()




#############################################################################

def doc_read_text(doc_path:str)->str:
    doc = docx.Document(doc_path)
    text = ''
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text += run.text
        text += '\n'
    return text
#############################################################################
def doc_read_text_highlight(doc_path:str)->str:
    doc = docx.Document(doc_path)
    text = ''


    for para_idx, paragraph in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(paragraph.runs):
            # 尝试获取文字的背景色（shd元素）
            shd = run._element.xpath('.//w:shd')
            if shd:
                fill = shd[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                if fill:
                    text += f'[HIGHLIGHT:{run.text}]\n'
    return text
#############################################################################
def doc_read_text_delete(doc_path:str)->str:
    doc = docx.Document(doc_path)
    text = ''

    for para_idx, paragraph in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(paragraph.runs):
            # 尝试获取文字的背景色（shd元素）
            xml_bytes = etree.tostring(run._element, encoding='unicode', pretty_print=True)
            # 打印格式化后的XML
            print(xml_bytes)
    return text
#############################################################################
def test_code():
    doc_path = 'highlight.docx'  # 替换为实际文档路径
    text = doc_read_text(doc_path)
    print(text)

    print("===============================")
    text = doc_read_text_highlight(doc_path)
    print(text)
    print("===============================")
    text = doc_read_text_delete(doc_path)
    print(text)